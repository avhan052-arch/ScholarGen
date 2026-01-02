import os
import google.oauth2.id_token
import google.auth.transport.requests
from fastapi import WebSocket, WebSocketDisconnect
from jose import jwt, JWTError
from typing import List, Dict
import requests
from datetime import datetime, timedelta
from fastapi import FastAPI, HTTPException, Request, status, Form, Depends, UploadFile, File
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse  # Import ini penting
from pydantic import BaseModel
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import create_engine, Column, Integer, String, Boolean, case, func
import re
import database
import auth
import time
from scholarly import scholarly
from database import TopUpRequest # Import model baru
from fastapi.staticfiles import StaticFiles
from auth import decode_access_token, SECRET_KEY, ALGORITHM
from bot import start_bot, notify_new_topup

# Masukkan Client ID Anda di sini
# GOOGLE_CLIENT_ID = "483904910670-di4quivnermuaa6utuq6qfv4dhhq5sol.apps.googleusercontent.com"
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "ID_LAMA_ANDA")

os.makedirs("uploads", exist_ok=True)

# 1. Setup App
app = FastAPI()
app.mount("/static_uploads", StaticFiles(directory="uploads"), name="uploads")
static_dir = os.path.join(os.path.dirname(__file__), "static")

app.mount("/static", StaticFiles(directory=static_dir), name="static")
# Izinkan Frontend mengakses Backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- CONNECTION MANAGER WEBSOCKET ---
class ConnectionManager:
    def __init__(self):
        # Dictionary untuk menyimpan koneksi user berdasarkan ID mereka
        # Contoh: { 1: [WebSocket1, WebSocket2], 2: [WebSocket3] }
        self.active_connections: Dict[int, List[WebSocket]] = {}

        # --- BARU: List khusus Admin (untuk update tabel) ---
        self.admin_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket, user_id: int):
        # await websocket.accept()
        if user_id not in self.active_connections:
            self.active_connections[user_id] = []
        self.active_connections[user_id].append(websocket)
        print(f"✅ User {user_id} connected via WebSocket.")

    def disconnect(self, websocket: WebSocket, user_id: int = None):
        if user_id is not None:
            if user_id in self.active_connections:
                if websocket in self.active_connections[user_id]:
                    self.active_connections[user_id].remove(websocket)
                # Jika tidak ada koneksi tersisa untuk user ini, hapus keynya
                if not self.active_connections[user_id]:
                    del self.active_connections[user_id]
            print(f"❌ User {user_id} disconnected.")
        else:
            # Jika tidak ada user_id, mungkin ini adalah admin websocket
            if websocket in self.admin_connections:
                self.admin_connections.remove(websocket)
            print(f"❌ Admin connection closed.")

        # Juga cek dan hapus dari admin connections jika ada
        if websocket in self.admin_connections:
            self.admin_connections.remove(websocket)
        print(f"❌ Connection closed.")

    # --- BARU: Broadcast ke SEMUA Admin ---
    async def broadcast_to_admins(self, message: dict):
        if self.admin_connections:
            for connection in self.admin_connections:
                try:
                    await connection.send_json(message)
                except:
                    pass # Jika admin offline/putus, abaikan

    async def broadcast_to_user(self, user_id: int, message: dict):
        if user_id in self.active_connections:
            for connection in self.active_connections[user_id]:
                try:
                    await connection.send_json(message)
                except:
                    pass # Jika koneksi putus, abaikan

manager = ConnectionManager()   

@app.websocket("/ws/admin")
async def ws_admin_endpoint(websocket: WebSocket):
    await websocket.accept()
    
    # 1. Ambil token
    token = websocket.query_params.get("token")
    
    if not token:
        print("⛔ WS Admin: Tidak ada token di URL.")
        await websocket.close(code=1008)
        return

    # 2. Validasi Langsung (Tanpa tergantung auth.py)
    try:
        # Coba decode token
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email = payload.get("sub")
        
        print(f"🛡️ WS Admin Connected! User: {email}")

    except JWTError as e:
        print(f"⛔ WS Admin Error JWT: {e}")
        print(f"⛔ Token yang diterima: {token[:20]}...") # Cetak potongan token
        await websocket.close(code=1008, reason="Invalid Token")
        return
    except Exception as e:
        print(f"⛔ WS Admin Error Lain: {e}")
        await websocket.close(code=1011, reason="Internal Error")
        return

    # 3. Hubungkan ke Manager
    manager.admin_connections.append(websocket)
    print(f"🛡️ Admin {email} connected to WebSocket.")

    try:
        while True:
            data = await websocket.receive_text()
    except WebSocketDisconnect:
        # Gunakan fungsi disconnect yang benar untuk admin
        if websocket in manager.admin_connections:
            manager.admin_connections.remove(websocket)
        print("🛡️ Admin Disconnected.")


@app.websocket("/ws/{user_id}")
async def websocket_endpoint(websocket: WebSocket, user_id: int):
    await websocket.accept()

    try:
        # Pastikan user_id adalah integer valid
    user_id = int(user_id)
    except ValueError:
        await websocket.close(code=1003)  # Invalid data
        return

    # Hubungkan ke manager
    await manager.connect(websocket, user_id)
    print(f"✅ User {user_id} connected to WebSocket")

    try:
        # Koneksi tetap terbuka (keep alive)
        while True:
            data = await websocket.receive_text()
            # Jika mau menerima pesan dari user, tangkap di sini
            # Untuk sekarang, kita hanya menerima pesan tanpa memproses
    except WebSocketDisconnect:
        manager.disconnect(websocket, user_id)
        print(f"❌ User {user_id} disconnected from WebSocket")

# --- PASTIKAN ADA KODE INI ---
# --- ENDPOINT: LOGIN GOOGLE ---
@app.post("/google-login")
async def google_login(token_data: dict, db: Session = Depends(auth.get_db)):
    """Endpoint ini menerima token dari Google, memverifikasinya,
    dan mengembalikan JWT aplikasi kita sendiri."""
    google_token = token_data.get("token")
    
    if not google_token:
        raise HTTPException(status_code=400, detail="Token Google tidak ditemukan")

    try:
        # 1. Verifikasi Token ke Server Google
        request = google.auth.transport.requests.Request()
        idinfo = google.oauth2.id_token.verify_oauth2_token(
            google_token, 
            request, 
            GOOGLE_CLIENT_ID
        )

        # 2. Ambil Email dari Token Google
        email = idinfo.get("email")
        if not email:
             raise HTTPException(status_code=400, detail="Gagal mengambil email dari Google")

        # 3. Cek User di Database
        user = db.query(database.User).filter(database.User.email == email).first()
        
        if not user:
            # 4. Jika User Belum Ada, Buatkan Otomatis (Auto-Register)
            # Pastikan password tidak melebihi 72 karakter untuk menghindari batasan bcrypt
            temp_password = "google_user_temp"[:72]
            random_password = auth.get_password_hash(temp_password)

            new_user = database.User(
                email=email,
                hashed_password=random_password,
                credits=3 # Saldo awal user baru
            )
            db.add(new_user)
            db.commit()
            db.refresh(new_user)
            user = new_user
            print(f"✅ Akun baru dibuat via Google: {email}")

            # --- TAMBAHKAN INI ---
            # Beritahu SEMUA ADMIN bahwa ada user baru terdaftar
            await manager.broadcast_to_admins({
                "type": "new_user",
                "data": {
                    "id": new_user.id,
                    "email": new_user.email,
                    "credits": new_user.credits,
                    "is_admin": new_user.is_admin
                }
            })
        else:
            print(f"✅ User lama login via Google: {email}")

        # 5. Generate JWT Aplikasi Kita
        access_token_expires = timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = auth.create_access_token(
            data={"sub": user.email}, 
            expires_delta=access_token_expires
        )

        return {
            "access_token": access_token, 
            "token_type": "bearer",
            "email": user.email,
            "credits": user.credits
        }

    # --- PERBAIKAN BAGIAN ERROR INI ---
    except ValueError as e:
        # Debug print ini sudah dirapikan indentasinya (4 spasi)
        print(f"❌ ERROR DETAIL GOOGLE LOGIN: {e}")
        raise HTTPException(status_code=400, detail=f"Token Google tidak valid: {e}")
    except Exception as e:
        print(f"❌ ERROR GENERAL GOOGLE LOGIN: {e}")
        # Cek apakah error terkait dengan batasan password bcrypt
        error_msg = str(e)
        if "password cannot be longer than 72 bytes" in error_msg:
            raise HTTPException(status_code=400, detail="Password terlalu panjang, maksimal 72 karakter")
        raise HTTPException(status_code=500, detail=f"Terjadi kesalahan server saat login Google: {e}")

# 1. SERVE FILE UPLOAD (AGAR ADMIN BISA LIHAT BUKTI TRANSFER)
# app.mount("/static_uploads", StaticFiles(directory="uploads"), name="uploads")


PRICING_MAP = {
    5: 15000,    # Paket 5 Kredit = Rp 15.000
    10: 35000,   # Paket 10 Kredit = Rp 35.000
    50: 150000,   # Paket 50 Kredit = Rp 150.000
    100: 300000,   # Paket 100 Kredit = Rp 100.000
    # Default fallback: 3000 per kredit jika user memasukkan jumlah lain
}

# Fungsi Helper untuk menambahkan kolom 'price' (Migrasi Database)
def migrate_database():
    try:
        from sqlalchemy import text
        engine = database.engine
        with engine.connect() as conn:
            # Coba tambahkan kolom price jika belum ada
            conn.execute(text("ALTER TABLE topup_requests ADD COLUMN price INTEGER DEFAULT 0"))
            conn.commit()
        print("✅ Kolom 'price' berhasil ditambahkan ke database.")
    except Exception as e:
        # Jika kolom sudah ada, akan muncul error, kita abaikan saja
        print(f"ℹ️ Info Migrasi: {e}")

# Ini memastikan tabel 'users' dibuat otomatis saat server nyala
@app.on_event("startup")
def startup_db_client():
    # Migrasi tabel
    migrate_database()

    # Buat tabel user
    database.Base.metadata.create_all(bind=database.engine)
    print("✅ Database Tables Checked/Created.")

    # Buat admin secara otomatis jika belum ada
    from auth import get_password_hash
    db = database.SessionLocal()

    # Ambil email dan password admin dari environment variables atau gunakan default
    admin_email = os.environ.get("ADMIN_EMAIL", "avhan43@gmail.com")
    admin_password = os.environ.get("ADMIN_PASSWORD", "admin123")

    # Cek apakah admin sudah ada
    existing_admin = db.query(database.User).filter(database.User.email == admin_email).first()

    if not existing_admin:
        # Buat admin baru
        hashed_password = get_password_hash(admin_password)
        new_admin = database.User(
            email=admin_email,
            hashed_password=hashed_password,
            is_admin=True,
            credits=99999  # Saldo awal untuk admin
        )
        db.add(new_admin)
        db.commit()
        print(f"✅ Admin baru dibuat otomatis: {admin_email}")
    else:
        # Jika user dengan email tersebut sudah ada, pastikan statusnya adalah admin dan password diupdate
        existing_admin.is_admin = True
        existing_admin.hashed_password = get_password_hash(admin_password)
        db.commit()
        print(f"✅ User {admin_email} diupdate menjadi Admin (password diupdate)")

    db.close()
    start_bot()


# 2. Setup Klien Groq
# Masukkan API Key Groq Anda di sini
def get_groq_client():
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY environment variable is not set")
    return Groq(api_key=api_key)

# --- 1. DEFINISIKAN STRUKTUR SKRIPSI YANG BAKU ---
PEDOMAN_SKRIPSI = {
    "Bab 1 Pendahuluan (Latar Belakang & Rumusan Masalah)": """
    TULISLAH DENGAN STRUKTUR BERIKUT WAJIB:
    1.1 Latar Belakang Masalah
       - Diawali dengan fenomena global/umum.
       - Mengerucut ke konteks nasional.
       - Mengerucut ke konteks spesifik/topik penelitian.
       - Akhiri dengan paragraf *ad absurdum* (ada yang belum terpecahkan).
    1.2 Rumusan Masalah
       - Berbentuk kalimat tanya yang jelas dan spesifik.
       - Minimal 2 pertanyaan terkait hubungan variabel.
    1.3 Tujuan Penelitian
       - Jawaban dari rumusan masalah (bukan kalimat tanya).
    1.4 Manfaat Penelitian
       - Manfaat Teoretis (untuk pengembangan ilmu).
       - Manfaat Praktis (untuk masyarakat/instansi).
    """,

    "Bab 2 Tinjauan Pustaka": """
    TULISLAH DENGAN STRUKTUR BERIKUT WAJIB:
    2.1 Landasan Teori
       - Definisi konsep variabel utama.
       - Teori-teori yang relevan dengan topik.
    2.2 Penelitian Terdahulu
       - Ringkasan 2-3 penelitian sebelumnya (gunakan referensi yang saya berikan).
       - Jelaskan kesamaan dan perbedaan dengan penelitian Anda.
    2.3 Kerangka Pemikiran
       - Gambaran hubungan antar variabel (dijelaskan secara narasi).
    2.4 Hipotesis
       - Dinyatakan dalam bentuk kalimat penegasan sementara.
    """,

    "Bab 3 Metodologi Penelitian": """
    TULISLAH DENGAN STRUKTUR BERIKUT WAJIB:
    3.1 Jenis Penelitian
       - Misal: Kuantitatif, Kualitatif, atau R&D.
       - Jelaskan alasannya.
    3.2 Lokasi dan Waktu Penelitian
       - Tempat dilaksanakannya penelitian.
    3.3 Populasi dan Sampel
       - Definisikan populasi (universe) dan teknik pengambilan sampel (Purposive/Random).
    3.4 Teknik Pengumpulan Data
       - Kuesioner, Wawancara, Observasi, atau Dokumentasi.
    3.5 Teknik Analisis Data
       - Uji regresi, uji t, deskriptif, dll (sesuaikan jenis penelitian).
    """,
    
    "Bab 4 Hasil dan Pembahasan": """
    TULISLAH DENGAN STRUKTUR BERIKUT WAJIB:
    4.1 Hasil Penelitian
       - Presentasi data yang didapat (deskriptif).
    4.2 Pembahasan
       - Menginterpretasikan hasil penelitian.
       - Menghubungkan hasil dengan teori di Bab 2.
       - Membahas apakah hipotesis terbukti atau tidak.
    """
}


# --- ADMIN DEPENDENCY ---
# Fungsi ini memastikan hanya user dengan is_admin=True yang bisa akses
async def get_current_admin_user(current_user: database.User = Depends(auth.get_current_user)):
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Anda tidak memiliki akses Admin.")
    return current_user

# --- ADMIN ENDPOINTS ---
@app.get("/admin", response_class=HTMLResponse)
async def get_admin_page():
    try:
        with open("admin.html", "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Error: File admin.html tidak ditemukan di folder ini!</h1>", status_code=404)

# 1. Route untuk menampilkan halaman login admin
@app.get("/admin-login", response_class=HTMLResponse)
async def get_admin_login_page():
    try:
        with open("admin_login.html", "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return HTMLResponse(content="<h1>File admin_login.html tidak ditemukan</h1>", status_code=404)

# 2. Endpoint Login Khusus Admin
@app.post("/admin/token")
async def admin_login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(auth.get_db)):
    try:
        # 1. Cari user berdasarkan email
        user = db.query(database.User).filter(database.User.email == form_data.username).first()

        # 2. Cek apakah user ada DAN password benar
        if not user or not auth.verify_password(form_data.password, user.hashed_password):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Email atau password salah"
            )

        # 3. CEK KRUSIAL: Apakah user ini ADMIN?
        if not user.is_admin:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Akses ditolak: Anda bukan Administrator"
            )

        # 4. Generate Token
        access_token_expires = timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = auth.create_access_token(
            data={"sub": user.email},
            expires_delta=access_token_expires
        )

        return {"access_token": access_token, "token_type": "bearer"}

    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        # Log error untuk debugging
        print(f"❌ Error di admin login: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat login"
        )
    finally:
        # Pastikan koneksi database ditutup
        db.close()

# 1. Ambil Semua Data User
@app.get("/admin/users")
async def get_all_users(admin: database.User = Depends(get_current_admin_user), db: Session = Depends(auth.get_db)):
    users = db.query(database.User).all()
    result = []
    for u in users:
        result.append({
            "id": u.id,
            "email": u.email,
            "credits": u.credits,
            "is_admin": u.is_admin,
            "created_at": str(u.created_at) if hasattr(u, 'created_at') else "-"
        })
    return {"users": result}


# Endpoint khusus untuk Dashboard Stats
@app.get("/admin/stats")
async def get_admin_stats(admin: database.User = Depends(get_current_admin_user), db: Session = Depends(auth.get_db)):

    try:
        # 1. Hitung Total User
        total_users = db.query(database.User).count()

        # 2. Hitung Request Pending
        pending_requests = db.query(database.TopUpRequest).filter(database.TopUpRequest.status == "Pending").count()

        # 3. Hitung Total Pendapatan (SUM 'price' WHERE status = Approved) <--- UBAH DISINI
        revenue_result = db.query(func.sum(database.TopUpRequest.price)).filter(database.TopUpRequest.status == "Approved").scalar()
        total_revenue = revenue_result if revenue_result else 0

        return {
            "total_users": total_users,
            "pending_requests": pending_requests,
            "total_revenue": total_revenue
        }
    except Exception as e:
        print(f"❌ Error di admin stats: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat mengambil statistik admin"
        )
    finally:
        db.close()

# 2. Ambil Semua Request Top Up (Global, bukan milik sendiri)
@app.get("/admin/topup_requests")
async def get_all_topups(admin: database.User = Depends(get_current_admin_user), db: Session = Depends(auth.get_db)):
    try:
        # requests = db.query(database.TopUpRequest).order_by(database.TopUpRequest.created_at.desc()).all()
        # Ambil data dengan logika urutan: Pending (Terbaru) -> Approved (Terbaru) -> Rejected (Terbaru)
        requests = db.query(TopUpRequest).order_by(
            # 1. Mengurutkan berdasarkan Status (Bobot Prioritas)
            case(
                (TopUpRequest.status == 'Pending', 0), # Paling Atas
                (TopUpRequest.status == 'Approved', 1),
                (TopUpRequest.status == 'Rejected', 2),
                else_=3
            ),
            # 2. Mengurutkan berdasarkan Waktu (Terbaru di dalem masing-masing status)
            TopUpRequest.created_at.desc()
        ).all()
        result = []
        for r in requests:
            user = db.query(database.User).filter(database.User.id == r.user_id).first()
            result.append({
                "id": r.id,
                "user_email": user.email if user else "Unknown",
                "amount": r.amount,
                "proof_filename": r.account_number, # Kita simpan nama file di sini
                "status": r.status,
                "created_at": str(r.created_at)
            })
            # Hapus print(result) karena ini akan mencetak data ke console
        return {"requests": result}
    except Exception as e:
        print(f"❌ Error di admin topup requests: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat mengambil request top up"
        )
    finally:
        db.close()

# 3. Aksi Admin: Setujui Top Up
@app.post("/admin/approve_topup/{request_id}")
async def admin_approve_topup(request_id: int, admin: database.User = Depends(get_current_admin_user), db: Session = Depends(auth.get_db)):
    try:
        req = db.query(database.TopUpRequest).filter(TopUpRequest.id == request_id).first()
        if not req: raise HTTPException(status_code=404, detail="Request tidak ditemukan")

        if req.status == "Approved":
            raise HTTPException(status_code=400, detail="Sudah disetujui sebelumnya")

        req.status = "Approved"

        # Tambah kredit ke user
        user = db.query(database.User).filter(database.User.id == req.user_id).first()
        user.credits += req.amount

        db.commit()
         # Ini penting untuk memastikan user.credits mencerminkan nilai di file DB
        db.refresh(user)

        revenue_result = db.query(func.sum(database.TopUpRequest.price)).filter(TopUpRequest.status == "Approved").scalar()
        current_total_revenue = revenue_result if revenue_result else 0

       # --- TAMBAHKAN INI ---
        # Beritahu User
        await manager.broadcast_to_user(user.id, {
            "type": "credit_update",
            "amount": user.credits,
            "message": f"Top up sebesar {req.amount} kredit disetujui!"
        })

        # Beritahu SEMUA ADMIN (Realtime Update Widget)
        await manager.broadcast_to_admins({
            "type": "update_revenue",
            "total_revenue": current_total_revenue
        })
        # -------------------
        return {"message": "Top Up Disetujui", "new_balance": user.credits}
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        print(f"❌ Error di admin approve topup: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat menyetujui top up"
        )
    finally:
        db.close()

# 4. Aksi Admin: Tolak Top Up
@app.post("/admin/reject_topup/{request_id}")
async def admin_reject_topup(request_id: int, admin: database.User = Depends(get_current_admin_user), db: Session = Depends(auth.get_db)):
    try:
        req = db.query(database.TopUpRequest).filter(TopUpRequest.id == request_id).first()
        if not req: raise HTTPException(status_code=404, detail="Request tidak ditemukan")

        req.status = "Rejected"
        db.commit()
         # --- TAMBAHKAN INI ---
        # Kirim notifikasi ke user yang request
        await manager.broadcast_to_user(req.user_id, {
            "type": "topup_notification",
            "status": "rejected",
            "message": "Maaf, permintaan Top Up Anda ditolak oleh Admin."
        })
        # -------------------

        return {"message": "Top Up Ditolak"}
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        print(f"❌ Error di admin reject topup: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat menolak top up"
        )
    finally:
        db.close()

# 5. Aksi Admin: Update Saldo Manual
@app.post("/admin/update_credits")
async def update_credits_manual(
    user_id: int = Form(...),
    amount: int = Form(...),
    action: str = Form(...), # 'add' or 'subtract'
    admin: database.User = Depends(get_current_admin_user),
    db: Session = Depends(auth.get_db)
):
    try:
        user = db.query(database.User).filter(database.User.id == user_id).first()
        if not user: raise HTTPException(status_code=404, detail="User tidak ditemukan")

        if action == 'add':
            user.credits += amount
        elif action == 'subtract':
            user.credits -= amount

        db.commit()
        await manager.broadcast_to_user(user_id, {
        "type": "credit_update",
        "amount": user.credits,
        "message": "Saldo diperbarui oleh Admin."
        })
        return {"message": "Saldo berhasil diupdate", "new_balance": user.credits}
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        print(f"❌ Error di admin update credits: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat memperbarui kredit"
        )
    finally:
        db.close()

# 6. Aksi Admin: Reset Kredit User
@app.post("/admin/reset_user/{user_id}")
async def reset_user_credits(
    user_id: int,
    admin: database.User = Depends(get_current_admin_user),
    db: Session = Depends(auth.get_db)
):
    try:
        user = db.query(database.User).filter(database.User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User tidak ditemukan")

        # Simpan nilai kredit sebelum direset untuk notifikasi
        old_credits = user.credits

        # Reset kredit ke nilai awal (3 kredit sesuai dengan nilai default saat registrasi)
        user.credits = 3

        db.commit()

        # Kirim notifikasi ke user yang bersangkutan
        await manager.broadcast_to_user(user_id, {
            "type": "credit_update",
            "amount": user.credits,
            "message": f"Kredit Anda telah direset oleh admin dari {old_credits} KR menjadi {user.credits} KR."
        })

        return {
            "message": "Kredit user berhasil direset ke nilai awal",
            "new_balance": user.credits,
            "old_balance": old_credits
        }
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        print(f"❌ Error di admin reset user: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat mereset kredit user"
        )
    finally:
        db.close()

# 7. Aksi Admin: Delete User
@app.delete("/admin/delete_user/{user_id}")
async def delete_user(
    user_id: int,
    admin: database.User = Depends(get_current_admin_user),
    db: Session = Depends(auth.get_db)
):
    try:
        user = db.query(database.User).filter(database.User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User tidak ditemukan")

        # Cek apakah user yang akan dihapus adalah admin
        if user.is_admin:
            raise HTTPException(status_code=400, detail="Tidak dapat menghapus admin")

        # Hapus user dari database
        db.delete(user)
        db.commit()

        return {"message": "User berhasil dihapus"}
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        print(f"❌ Error di admin delete user: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Terjadi kesalahan internal saat menghapus user"
        )
    finally:
        db.close()

# Fungsi untuk mengambil struktur berdasarkan input user
# Helper: Struktur Skripsi Baku
def get_structured_instruction(jenis, bab):
    if jenis in ["Skripsi", "Tesis"]:
        struktur = {
            "Bab 1 Pendahuluan": """
            1.1 Latar Belakang Masalah
            1.2 Rumusan Masalah
            1.3 Tujuan Penelitian
            1.4 Manfaat Penelitian
            """,
            "Bab 2 Tinjauan Pustaka": """
            2.1 Landasan Teori
            2.2 Penelitian Terdahulu
            2.3 Kerangka Pemikiran
            2.4 Hipotesis
            """,
            "Bab 3 Metodologi": """
            3.1 Jenis Penelitian
            3.2 Populasi dan Sampel
            3.3 Teknik Pengumpulan Data
            3.4 Teknik Analisis Data
            """,
            "Bab 4 Hasil dan Pembahasan": """
            4.1 Hasil Penelitian
            4.2 Pembahasan
            """,
            "Bab 5 Kesimpulan": """
            5.1 Kesimpulan
            5.2 Saran
            """,
            "Daftar Pustaka": """
            """
        }
        return struktur.get(bab, "Tulislah secara akademis dengan struktur logis.")
    return "Tulislah secara akademis."


# Helper: Google Scholar Scraper
def search_google_scholar(query, limit=3):
    results = []
    try:
        # Mencari query
        search_query = scholarly.search_pubs(query)
        
        # Ambil limit hasil (hanya 3-5 biar cepat)
        for i in range(limit):
            try:
                pub = next(search_query)
                
                # Parsing data dari format Google Scholar
                # Data Google Scholar berbeda dengan OpenAlex, kita rapikan jadi sama
                pub_info = pub.get('bib', {})
                authors_str = " & ".join(pub_info.get('author', [])) # Daftar penulis digabung &
                
                results.append({
                    "title": pub_info.get('title', 'Tanpa Judul'),
                    "authors": authors_str,
                    "year": pub_info.get('pub_year', 'Unknown'),
                    "venue": pub_info.get('journal', 'Google Scholar'), 
                    "citation_count": pub.get('num_citations', 0),
                    "citation_apa": f"{authors_str} ({pub_info.get('pub_year')}). {pub_info.get('title')}. {pub_info.get('journal', 'Google Scholar')}.",
                    "pdf_url": pub.get('eprint_url'), # Link PDF jika ada
                    "source": "Google Scholar" # Penanda agar tahu asalnya
                })
            except StopIteration:
                break # Selesai jika data habis
            except Exception as e:
                print(f"Error parsing satu hasil GS: {e}")
                continue
                
    except Exception as e:
        print(f"⚠️ Google Scholar Error (Biasanya diblokir IP): {e}")
        # Jika error, kita return list kosong (abaikan), jangan crash aplikasi
    
    return results

class GenerationRequest(BaseModel):
    jenis_dokumen: str
    judul: str
    bab: str
    kata_kunci: str = ""
    selected_references: List[str] = []

# --- TAMBAHKAN FUNGSI INI KE main.py ---

# --- FITUR TOP UP KREDIT ---
# --- FITUR REQUEST TOP UP MANUAL ---
@app.post("/request_topup")
async def request_topup(
    amount: int = Form(...),           # Menerima data form biasa
    proof: UploadFile = File(...),     # Menerima file gambar
    current_user: database.User = Depends(auth.get_current_user)
):
    db = database.SessionLocal()
    
    # 1. Simpan File Bukti ke Folder Server
    try:
        # --- 1. HITUNG HARGA AKHIR ---
        final_price = PRICING_MAP.get(amount, amount * 3000)
        # Buat folder 'uploads' jika belum ada
        # os.makedirs("uploads", exist_ok=True)
        
        # Nama file = UserID + Timestamp + Ekstensi asli
        file_location = f"uploads/{current_user.id}_{int(time.time())}_{proof.filename}"
        filename_to_save = f"{current_user.id}_{int(time.time())}_{proof.filename}"
        
        with open(file_location, "wb+") as file_object:
            file_object.write(await proof.read()) # Tulis file ke disk
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal upload file: {str(e)}")

    # 2. Simpan Data ke Database
    try:
        new_request = database.TopUpRequest(
            user_id = current_user.id,
            amount = amount,
            method = "Bank Transfer", 
            account_number = filename_to_save, # Simpan nama file di kolom account_number
            status = "Pending",
            price = final_price
        )
        db.add(new_request)
        db.commit()
        db.refresh(new_request)
        
        # --- TAMBAHKAN INI (Sebelum return) ---
        # Beritahu semua admin bahwa ada request baru
        await manager.broadcast_to_admins({
            "type": "new_request",
            "data": {
                "id": new_request.id,
                "created_at": str(new_request.created_at),
                "user_email": current_user.email,
                "amount": new_request.amount,
                "status": new_request.status,
                "proof_filename": new_request.account_number
            }
        })
        # --- KIRIM NOTIFIKASI TELEGRAM ---
        # Ambil nama user (current_user.email)
        notify_new_topup(
            amount=new_request.amount,
            user_email=current_user.email,
            image_filename=new_request.account_number, # Nama file gambar
            request_id=new_request.id
        )
        # ------------------------------

        return {
            "message": "Request dan Bukti Berhasil Dikirim! Menunggu konfirmasi admin.",
            "status": "Pending"
        }

        
    except Exception as e:
        # Jika simpan DB gagal, hapus file yang sudah diupload (cleanup)
        if os.path.exists(file_location):
            os.remove(file_location)
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()

# --- (Opsional) ENDPOINT ADMIN: Setujui Request ---
# Ini untuk simulasi Anda seolah-olah menyetujui user secara manual
@app.post("/admin_approve_topup/{request_id}")
async def admin_approve(request_id: int, db: database.SessionLocal = Depends(database.get_db)):
    req = db.query(TopUpRequest).filter(TopUpRequest.id == request_id).first()
    if not req: raise HTTPException(status_code=404, detail="Request tidak ada")
    
    req.status = "Approved"
    
    # Tambahkan kredit ke user pemilik request
    user = db.query(database.User).filter(database.User.id == req.user_id).first()
    user.credits += req.amount
    
    db.commit()
    # --- TAMBAHKAN INI ---
    await manager.broadcast_to_admins({
        "type": "update_request",
        "id": req.id,
        "new_status": "Approved"
    })

    return {"message": "Top Up Disetujui dan Kredit Ditambahkan"}

# --- Endpoint: Lihat Riwayat Top Up ---
@app.get("/my_topups")
async def get_my_topups(current_user: database.User = Depends(auth.get_current_user)):
    db = database.SessionLocal()
    try:
        requests = db.query(TopUpRequest).filter(TopUpRequest.user_id == current_user.id).all()
        
        # Format data untuk frontend
        history = []
        for r in requests:
            history.append({
                "id": r.id,
                "amount": r.amount,
                "method": r.method,
                "status": r.status,
                "created_at": str(r.created_at)
            })
        
        return {"history": history}
    finally:
        db.close()


@app.get("/me")
async def get_me(current_user: database.User = Depends(auth.get_current_user)):
    # Buat sesi DB baru untuk mengambil data terbaru
    # (Kita tidak menggunakan current_user langsung karena mungkin objeknya lama)
    db = database.SessionLocal()
    
    # Ambil ulang data user dari database menggunakan ID dari token
    fresh_user = db.query(database.User).filter(database.User.id == current_user.id).first()
    
    db.close()

    if not fresh_user:
        return {"error": "User not found"}

    return {
        "id": fresh_user.id,
        "email": fresh_user.email,
        "credits": fresh_user.credits
    }

@app.post("/search")
async def search_references(request: dict, current_user: database.User = Depends(auth.get_current_user)):
    query = request.get("query", "")
    if not query: raise HTTPException(status_code=400, detail="Query kosong")

    all_results = []
    
    # 1. CARI DI OPENALEX (Sumber Utama)
    headers = {'User-Agent': 'Mozilla/5.0 ScholarGenApp/1.0', 'Accept': 'application/json'}
    url = f"https://api.openalex.org/works?search={query}&per_page=3&sort=relevance_score:desc"
    
    try:
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            data = response.json()
            for work in data.get('results', []):
                authorships = work.get('authorships', [])
                authors = ", ".join([a['author']['display_name'] for a in authorships[:3]])
                all_results.append({
                    "title": work['title'],
                    "authors": authors,
                    "year": work.get('publication_year'),
                    "venue": work.get('primary_location', {}).get('source', {}).get('display_name', 'Unknown'),
                    "citation_count": work.get('cited_by_count', 0),
                    "citation_apa": f"{authors} ({work.get('publication_year')}). {work['title']}. {work.get('primary_location', {}).get('source', {}).get('display_name', 'Unknown')}.",
                    "pdf_url": work.get('best_oa_location', {}).get('pdf_url'),
                    "source": "OpenAlex"
                })
    except Exception as e:
        # JANGAN GUNAKAN 'pass' LAGI. Log error-nya agar kita tahu penyebabnya.
        print(f"⚠️ Gagal menghubungi OpenAlex: {e}")


    # 2. CARI DI GOOGLE SCHOLAR (Sumber Tambahan)
    try:
        gs_results = search_google_scholar(query, limit=5)
        # Gabungkan hasil
        all_results.extend(gs_results)
    except: pass

    # 3. Return Semua Hasil
    return {"status": "success", "data": all_results}

        
# --- BARU: Route untuk menampilkan Halaman Utama ---
@app.get("/", response_class=HTMLResponse)
async def read_root():
    try:
        with open("landing.html", "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Error: File landing.html tidak ditemukan!</h1>", status_code=404)

# --- ROUTE DASHBOARD: APLIKASI UTAMA ---
@app.get("/app", response_class=HTMLResponse)
async def read_app():
    try:
        with open("app.html", "r", encoding="utf-8") as f: # Pastikan nama file ini sesuai rename di langkah 1
            return f.read()
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Error: File app.html tidak ditemukan!</h1>", status_code=404)


# Fungsi Prompt Engineering
def create_system_prompt(jenis, bab):
    base_instruction = "Anda adalah Asisten Penulis Akademik Profesional. Tugas Anda adalah menulis konten akademis yang formal, terstruktur, dan ilmiah."
    
    specific_instruction = f"""
    Tulis bagian '{bab}' untuk {jenis}.
    
    Aturan Penulisan:
    1. Gunakan bahasa Indonesia formal.
    2. Struktur teks dengan heading yang jelas (H1, H2, H3) menggunakan format Markdown.
    3. Sertakan penomoran poin (1.1, 1.2, dst).
    4. Jangan membuat referensi fiktif. Gunakan placeholder [Penulis, Tahun] jika perlu menyebut studi.
    5. Fokus pada substansi dan kedalaman analisis.
    """
    return f"{base_instruction}\n{specific_instruction}"

# --- FITUR BARU: EXPORT TO WORD ---
@app.post("/export-word")
async def export_to_word(request: dict):
    """
    Menerima text dari frontend, mengonversi ke format Word, dan mendownloadnya.
    Body JSON: {"content": "isi tulisan...", "filename": "nama_file"}
    """
    content = request.get("content", "")
    filename = request.get("filename", "Draft_Akademik")

    # 1. Buat Dokumen Word Baru
    doc = Document()
    
    # Set margin standar skripsi (3-4 cm) - Opsional, di sini kita pakai default dulu
    section = doc.sections[0]
    section.top_margin = Inches(1)     # 1 inci
    section.bottom_margin = Inches(1)   # 1 inci
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

    # 2. Parsing Teks (Konversi Markdown Sederhana ke Word Style)
    # Kita split per baris untuk mendeteksi Heading
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Tambahkan paragraf kosong
            doc.add_paragraph()
            continue
        
        # Deteksi Heading 1 (Markdown: # Judul)
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            
        # Deteksi Heading 2 (Markdown: ## Judul)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            
        # Deteksi Heading 3 (Markdown: ### Judul)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            
        # Deteksi List (Markdown: - item)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
            
        # Paragraf Biasa (dengan dukungan Bold sederhana)
        else:
            p = doc.add_paragraph()
            # Kita cek format bold menggunakan regex sederhana **teks**
            # Ini cara kasar tapi cukup untuk MVP
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Jika ini teks bold, tambahkan run bold
                    run = p.add_run(part[2:-2]) # Hapus tanda **
                    run.bold = True
                else:
                    p.add_run(part)
    
    # 3. Simpan File Sementara
    output_path = f"{filename}.docx"
    doc.save(output_path)
    
    # 4. Kirim File ke Browser
    return FileResponse(
        output_path, 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=output_path
    )


@app.post("/export-complete")
async def export_complete_document(request: dict):
    """
    Menerima JSON: { "title": "...", "chapters": {"Bab 1": "text...", "Bab 2": "text..."} }
    """
    title = request.get("title", "Skripsi")
    chapters = request.get("chapters") # Ini adalah dictionary/object

    doc = Document()
    
    # Set Margin Halaman (Biasanya skripsi 3-4 cm, sini pakai 1.5 inch demo)
    section = doc.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    # 1. Halaman Judul (Opsional)
    doc.add_heading(title.upper(), 0)
    doc.add_paragraph("Dokumen ini digenerate otomatis oleh AI.")
    doc.add_paragraph("\n")

    # 2. Loop setiap bab yang dikirim dari frontend
    # Kita urutkan berdasarkan key (Bab 1, Bab 2...) agar berurutan
    sorted_chapters = sorted(chapters.items()) 

    for chapter_title, content in sorted_chapters:
        # Tambah Page Break sebelum bab baru (kecuali bab pertama mungkin)
        if chapter_title != sorted_chapters[0][0]:
            doc.add_page_break()
        
        # Tambah Judul Bab
        doc.add_heading(chapter_title, level=1)
        
        # Proses Isi Bab (sama seperti parsing sebelumnya)
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            # Parsing sederhana (Heading H2, H3, Bold)
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    # Simpan dan Kirim
    output_filename = f"{title}_Lengkap.docx"
    doc.save(output_filename)

    return FileResponse(
        output_filename, 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=output_filename
    )

# Endpoint Utama Generator
@app.post("/generate")
async def generate_text(
    request: GenerationRequest, 
    current_user: database.User = Depends(auth.get_current_user)
):
    # 1. Cek Kredit SaaS
    if current_user.credits <= 0:
        raise HTTPException(status_code=402, detail="Kredit habis! Silakan top up.")

    try:
        #---------------------------------------------------------
        # LANGKAH 1: PAKAI REFERENSI YANG DIPILIH USER (WAJIB)
        # ---------------------------------------------------------
        print(f"✅ User memilih {len(request.selected_references)} referensi.")
        
        # Ubah instruksi menjadi sangat tegas
        references_context = "PERINTAH KHUSUS: ANDA WAJIB MENGGUNAKAN REFERENSI BERIKUT INI SEBAGAI SATU-SATUNYA SUMBER DATA. JANGAN MENGGUNAKAN INFORMASI DI LUAR DAFTAR INI.\n\n"
        
        # Inisialisasi variable ini agar tidak error
        found_references = [] 

        for ref in request.selected_references:
            references_context += f"- {ref}\n"
            
            # Format ulang menjadi dictionary agar Frontend bisa baca
            found_references.append({
                "title": ref, 
                "authors": "User Selected",
                "year": "2024"
            })
        
        # Jika user tidak memilih apa-apa, gunakan teks umum
        if not request.selected_references:
            references_context = "Gunakan pengetahuan umum."
            found_references = []
        
        # ---------------------------------------------------------
        # LANGKAH 2: BUAT PROMPT UNTUK AI (TANPA OPSIONAL)
        # ---------------------------------------------------------
        struktur_baku = get_structured_instruction(request.jenis_dokumen, request.bab)
        
        system_prompt = f"""
        Anda adalah Penulis Akademik Profesional bidang {request.jenis_dokumen}.
        
        TUGAS ANDA:
        Tulis draft '{request.bab}' untuk topik: '{request.judul}'.
        
        ATURAN WAJIB (STANDAR PTKPT):
        1. Ikuti struktur sub-bab berikut TEPAT SEPERTI TERTULIS:
           {struktur_baku}
           
        2. GAYA BAHASA:
           - Gunakan Bahasa Indonesia baku sesuai EYD.
           - Kalimat efektif, tidak bertele-tele, hindari kata-kata slang.
           - Gunakan bahasa ilmiah (formal).
        
        3. KONTEN (PENTING):
           - {references_context}
           - WAJIB mensitasi referensi yang diberikan dengan format (Penulis, Tahun).
           - Untuk BAB 1 Tidak perlu langsung mensitasi nya.
           - untuk BAB 4  HASIL DAN PEMBAHASAN tidak perlu menyertakan sitasi cukup hasil dan pembahasan
           - Daftar Pustaka gunakan format sitasi APA, daftar pustakan berada di bab terakhir setelah Bab 4 Hasil dan Pembahasan
           - DILARANG KERAS membuat referensi fiktif/halusinasi di luar daftar ini.
        """
        
        user_prompt = f"Topik: {request.judul}. Kata Kunci: {request.kata_kunci}. Tulis sekarang."

        # PANGGIL API GROQ
        # Catatan: Saya ubah model ke 'llama3-70b-8192' agar stabil. 
        # Jika ingin memakai yang Anda tulis, pastikan nama modelnya valid di Groq.
        groq_client = get_groq_client()
        completion = groq_client.chat.completions.create(
            # model="llama-3.3-70b-versatile",
            model="openai/gpt-oss-120b",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.6,
            max_tokens=2048,
        )

        generated_content = completion.choices[0].message.content
        
        # ---------------------------------------------------------
        # LANGKAH 3: UPDATE DATABASE (KURANGI KREDIT)
        # ---------------------------------------------------------
        # Hanya eksekusi bagian ini jika AI sukses
        db = database.SessionLocal() # Buat sesi DB baru untuk update
        current_user.credits -= 1
        db.merge(current_user)
        db.commit()

        # --- TAMBAHKAN INI ---
        await manager.broadcast_to_user(current_user.id, {
            "type": "credit_update", 
            "amount": current_user.credits
        })
        db.close()
        
        print(f"✅ Kredit user {current_user.email} dikurangi. Sisa: {current_user.credits}")

        # ---------------------------------------------------------
        # LANGKAH 4: KEMBALIKAN HASIL
        # ---------------------------------------------------------
        return {
            "status": "success", 
            "data": generated_content,
            "sources": found_references,
            "remaining_credits": current_user.credits
        }

    except HTTPException:
        # Jangan tangkap HTTPException, biarkan naik
        raise
    except Exception as e:
        print(f"❌ Error Generate: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class RefineRequest(BaseModel):
    bab_key: str       # Kunci bab (misal: "Bab 1 Pendahuluan")
    current_content: str # Teks markdown asli yang sudah tersimpan
    instruction: str     # Perintah revisi dari user

@app.post("/refine")
async def refine_chapter(
    request: RefineRequest, 
    current_user: database.User = Depends(auth.get_current_user)
):
    # 1. Cek Kredit
    if current_user.credits <= 0:
        raise HTTPException(status_code=402, detail="Kredit habis! Edit membutuhkan kredit.")
    
    try:
        # 2. Prompt untuk AI Editor
        # Perintah: Tolong revisi atau tambahkan konten berdasarkan instruksi user
        system_prompt = f"""
        Anda adalah Editor Akademik yang canggih.
        Tugas Anda adalah MEREVISI/MENAMBAHKAN konten pada teks yang sudah ada.
        
        ATURAN:
        1. Jangan mengubah keseluruhan struktur secara drastis kecuali diminta.
        2. Jika diminta "Tambahkan", lakukan itu dengan mulus.
        3. Pertahankan gaya bahasa akademis.
        4. Pertahankan format Markdown (Heading ##, dll).
        """
        
        user_prompt = f"""
        INSTRUKSI USER: {request.instruction}
        
        TEKS SAAT INI (Markdown):
        {request.current_content}
        
        Tulis ulang teks tersebut sesuai instruksi user (dengan penambahan/revisi):
        """

        # Panggil AI
        groq_client = get_groq_client()
        completion = groq_client.chat.completions.create(
            # model="llama-3.3-70b-versatile",
            model="openai/gpt-oss-120b",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.6,
            max_tokens=2048,
        )

        revised_content = completion.choices[0].message.content

        # 3. Kurangi Kredit
        db = database.SessionLocal()
        current_user.credits -= 1
        db.merge(current_user)
        db.commit()
        db.close()

        return {
            "status": "success",
            "data": revised_content,
            "remaining_credits": current_user.credits
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
